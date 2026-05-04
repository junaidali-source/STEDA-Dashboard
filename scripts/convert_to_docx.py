"""
Convert Rumi_Causation_Experiment_Design.md to a formatted .docx file.
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colors ────────────────────────────────────────────────────────────────────
BLUE       = RGBColor(0x1F, 0x77, 0xB4)
ORANGE     = RGBColor(0xE6, 0x7E, 0x22)
GREEN      = RGBColor(0x27, 0xAE, 0x60)
DARK       = RGBColor(0x2C, 0x3E, 0x50)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xF4, 0xF6, 0xF7)
MID_GREY   = RGBColor(0xBD, 0xC3, 0xC7)
CODE_BG    = RGBColor(0xEC, 0xF0, 0xF1)
CODE_FG    = RGBColor(0x2C, 0x3E, 0x50)

# ── Helpers ───────────────────────────────────────────────────────────────────

def rgb_hex(rgb: RGBColor) -> str:
    return '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(rgb))
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   kwargs.get('val',   'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz',    '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', 'BDBDBD'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_run_with_style(para, text, bold=False, italic=False,
                       color=None, size=None, font='Calibri', code=False):
    run = para.add_run(text)
    run.font.name = 'Courier New' if code else font
    if size:  run.font.size = Pt(size)
    if bold:  run.font.bold = True
    if italic: run.font.italic = True
    if color: run.font.color.rgb = color
    if code:
        run.font.color.rgb = CODE_FG
        run.font.size = Pt(9)
    return run

def parse_inline(para, text, base_color=DARK, base_size=10.5):
    """Parse **bold**, *italic*, `code` inline markers and add runs."""
    pattern = re.compile(r'(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)')
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            add_run_with_style(para, part[2:-2], bold=True, color=base_color, size=base_size)
        elif part.startswith('*') and part.endswith('*'):
            add_run_with_style(para, part[1:-1], italic=True, color=base_color, size=base_size)
        elif part.startswith('`') and part.endswith('`'):
            add_run_with_style(para, part[1:-1], code=True)
        else:
            add_run_with_style(para, part, color=base_color, size=base_size)

def set_para_spacing(para, before=0, after=4):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)

def add_section_heading(doc, text, level=1):
    """Styled section heading with colored left bar via shading."""
    if level == 1:
        p = doc.add_paragraph()
        set_para_spacing(p, before=16, after=6)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  '1F77B4')
        pPr.append(shd)
        p.paragraph_format.left_indent  = Inches(0.15)
        p.paragraph_format.right_indent = Inches(0.0)
        run = p.add_run('  ' + text)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(14)
        run.font.bold  = True
        run.font.color.rgb = WHITE
        return p
    elif level == 2:
        p = doc.add_paragraph()
        set_para_spacing(p, before=10, after=3)
        run = p.add_run(text)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(12)
        run.font.bold  = True
        run.font.color.rgb = BLUE
        # bottom border
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '4')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), '1F77B4')
        pBdr.append(bot)
        pPr.append(pBdr)
        return p
    else:
        p = doc.add_paragraph()
        set_para_spacing(p, before=6, after=2)
        run = p.add_run(text)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(10.5)
        run.font.bold  = True
        run.font.color.rgb = ORANGE
        return p

def add_body(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=0, after=4)
    p.paragraph_format.left_indent = Inches(0.0)
    parse_inline(p, text.strip())
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    set_para_spacing(p, before=0, after=2)
    indent = Inches(0.25 * (level + 1))
    p.paragraph_format.left_indent  = indent
    p.paragraph_format.first_line_indent = Inches(-0.18)
    parse_inline(p, text.strip())
    return p

def add_numbered(doc, text, num):
    p = doc.add_paragraph()
    set_para_spacing(p, before=0, after=2)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    parse_inline(p, f"{num}. {text.strip()}")
    return p

def add_quote(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=4, after=4)
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'EBF5FB')
    pPr.append(shd)
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '12')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), '1F77B4')
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text.strip().lstrip('> ').strip('*'))
    run.font.italic = True
    run.font.size   = Pt(10)
    run.font.color.rgb = DARK
    return p

def add_code_block(doc, lines):
    p = doc.add_paragraph()
    set_para_spacing(p, before=4, after=4)
    p.paragraph_format.left_indent = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'ECF0F1')
    pPr.append(shd)
    code_text = '\n'.join(lines)
    run = p.add_run(code_text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(8.5)
    run.font.color.rgb = CODE_FG
    return p

def parse_md_table(lines):
    """Return (headers, rows) from markdown table lines."""
    rows = []
    for line in lines:
        if re.match(r'^\s*\|[-: |]+\|\s*$', line):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return [], []
    return rows[0], rows[1:]

def add_md_table(doc, headers, rows):
    col_count = max(len(headers), max((len(r) for r in rows), default=0))
    if col_count == 0:
        return
    # auto column widths
    page_width = Inches(6.3)
    col_w = page_width / col_count

    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers[:col_count]):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, BLUE)
        set_cell_border(cell, color='1F77B4')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        # strip bold markers from header
        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', h)
        run = p.add_run(clean)
        run.font.bold  = True
        run.font.size  = Pt(9.5)
        run.font.color.rgb = WHITE
        run.font.name  = 'Calibri'

    # data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = LIGHT_GREY if ri % 2 == 0 else RGBColor(0xFF, 0xFF, 0xFF)
        for ci in range(col_count):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            set_cell_border(cell, color='BDC3C7')
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            text = row[ci] if ci < len(row) else ''
            parse_inline(p, text, base_size=9)

    doc.add_paragraph()  # spacing after table

# ── Build Document ─────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.font.color.rgb = DARK

# ── Title Page ─────────────────────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(title_para, before=24, after=4)
r = title_para.add_run('Rumi Platform')
r.font.name  = 'Calibri'
r.font.size  = Pt(26)
r.font.bold  = True
r.font.color.rgb = BLUE

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(sub_para, before=0, after=4)
r = sub_para.add_run('Randomized Controlled Experiment Design')
r.font.name  = 'Calibri'
r.font.size  = Pt(18)
r.font.bold  = True
r.font.color.rgb = DARK

ital_para = doc.add_paragraph()
ital_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(ital_para, before=0, after=16)
r = ital_para.add_run('"Does Video Generation Cause Retention, or Do High-Intent Users Self-Select?"')
r.font.name   = 'Calibri'
r.font.size   = Pt(12)
r.font.italic = True
r.font.color.rgb = RGBColor(0x55, 0x6B, 0x82)

# Meta box
meta_table = doc.add_table(rows=4, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ('Document Type',   'Experiment Design / Research Protocol'),
    ('Date',            'April 28, 2026'),
    ('Based On',        'Video_Retention_Analysis_Report.pdf — Correlation Analysis'),
    ('Primary Question','Does introducing video generation cause users to stay longer, or does the correlation exist because high-intent users self-select?'),
]
for i, (label, value) in enumerate(meta_data):
    lc = meta_table.rows[i].cells[0]
    vc = meta_table.rows[i].cells[1]
    set_cell_bg(lc, LIGHT_GREY)
    set_cell_bg(vc, RGBColor(0xFF, 0xFF, 0xFF))
    set_cell_border(lc, color='BDC3C7')
    set_cell_border(vc, color='BDC3C7')
    lp = lc.paragraphs[0]
    lp.paragraph_format.space_before = Pt(3)
    lp.paragraph_format.space_after  = Pt(3)
    lr = lp.add_run(label)
    lr.font.bold = True; lr.font.size = Pt(9.5); lr.font.color.rgb = BLUE
    vp = vc.paragraphs[0]
    vp.paragraph_format.space_before = Pt(3)
    vp.paragraph_format.space_after  = Pt(3)
    vr = vp.add_run(value)
    vr.font.size = Pt(9.5); vr.font.color.rgb = DARK

doc.add_paragraph()
doc.add_page_break()

# ── Parse and render Markdown ──────────────────────────────────────────────────

MD_FILE = 'Rumi_Causation_Experiment_Design.md'
with open(MD_FILE, encoding='utf-8') as f:
    raw_lines = f.readlines()

lines = [l.rstrip('\n') for l in raw_lines]

i = 0
numbered_counter = 0
while i < len(lines):
    line = lines[i]

    # skip title lines (already on cover page)
    if i < 2 and line.startswith('#'):
        i += 1
        continue

    # horizontal rule
    if re.match(r'^---+$', line.strip()):
        p = doc.add_paragraph()
        set_para_spacing(p, before=6, after=6)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '4')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), 'BDC3C7')
        pBdr.append(bot)
        pPr.append(pBdr)
        i += 1
        continue

    # H1
    if line.startswith('# ') and not line.startswith('## '):
        add_section_heading(doc, line[2:].strip(), level=1)
        numbered_counter = 0
        i += 1
        continue

    # H2
    if line.startswith('## ') and not line.startswith('### '):
        add_section_heading(doc, line[3:].strip(), level=1)
        numbered_counter = 0
        i += 1
        continue

    # H3
    if line.startswith('### '):
        add_section_heading(doc, line[4:].strip(), level=2)
        numbered_counter = 0
        i += 1
        continue

    # H4
    if line.startswith('#### '):
        add_section_heading(doc, line[5:].strip(), level=3)
        i += 1
        continue

    # ARM labels (### ARM A:  etc) already caught by H3

    # code block
    if line.strip().startswith('```'):
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        add_code_block(doc, code_lines)
        i += 1
        continue

    # markdown table — collect all consecutive table lines
    if line.strip().startswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i])
            i += 1
        headers, rows = parse_md_table(table_lines)
        if headers:
            add_md_table(doc, headers, rows)
        continue

    # blockquote
    if line.strip().startswith('>'):
        add_quote(doc, line.strip())
        i += 1
        continue

    # numbered list
    m = re.match(r'^(\s*)\d+\.\s+(.*)', line)
    if m:
        numbered_counter += 1
        add_numbered(doc, m.group(2), numbered_counter)
        i += 1
        continue

    # bullet list (- or *)
    m = re.match(r'^(\s*)[-*]\s+(.*)', line)
    if m:
        numbered_counter = 0
        indent_level = len(m.group(1)) // 2
        add_bullet(doc, m.group(2), level=indent_level)
        i += 1
        continue

    # checkbox list
    m = re.match(r'^\s*-\s+\[[ x]\]\s+(.*)', line)
    if m:
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=2)
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run('[ ]  ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = MID_GREY
        parse_inline(p, m.group(1))
        i += 1
        continue

    # blank line
    if not line.strip():
        numbered_counter = 0
        i += 1
        continue

    # italic metadata footer
    if line.strip().startswith('*') and line.strip().endswith('*'):
        p = doc.add_paragraph()
        set_para_spacing(p, before=4, after=2)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line.strip().strip('*'))
        r.font.italic = True
        r.font.size   = Pt(9)
        r.font.color.rgb = MID_GREY
        i += 1
        continue

    # regular paragraph
    numbered_counter = 0
    add_body(doc, line)
    i += 1

# ── Save ───────────────────────────────────────────────────────────────────────
OUT = 'Rumi_Causation_Experiment_Design.docx'
doc.save(OUT)
print(f'[OK] Saved: {OUT}')

import os
print(f'     Size: {os.path.getsize(OUT)/1024:.1f} KB')
